from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "draft.md"
OUTPUT = ROOT / "manuscript" / "submission-version.docx"

EAST_ASIA_FONT = "SimSun"
LATIN_FONT = "Calibri"


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = LATIN_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.333, first_line=True, justify=True):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Pt(22)
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_style_font(style, size_pt: float, color: str = "000000", bold: bool = False):
    font = style.font
    font.name = LATIN_FONT
    font.size = Pt(size_pt)
    font.bold = bold
    font.color.rgb = RGBColor.from_string(color)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 11, "000000")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.333
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    set_run_font(run, 9, color="666666")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    page_run = OxmlElement("w:r")
    page_text = OxmlElement("w:t")
    page_text.text = "1"
    page_run.append(page_text)
    field.append(page_run)
    footer._p.append(field)
    run = footer.add_run(" 页")
    set_run_font(run, 9, color="666666")


def strip_inline_citations(text: str) -> str:
    return text


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    run = p.add_run(strip_inline_citations(text))
    set_run_font(run, 11)
    return p


def add_reference_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.167
    fmt.left_indent = Cm(0.74)
    fmt.first_line_indent = Cm(-0.74)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, 10)


def is_markdown_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def add_markdown_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return

    col_count = len(rows[0])
    table = doc.add_table(rows=0, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table)

    widths_by_count = {
        3: [2200, 3400, 3760],
        4: [1800, 2400, 2600, 2560],
    }
    widths = widths_by_count.get(col_count, [int(9360 / col_count)] * col_count)

    for row_index, row_cells in enumerate(rows):
        cells = table.add_row().cells
        for col_index, text in enumerate(row_cells):
            cell = cells[col_index]
            cell.width = Pt(widths[col_index] / 20)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, "F4F6F9")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.167
            run = paragraph.add_run(text)
            set_run_font(run, 10, bold=(row_index == 0))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def parse_markdown(lines: list[str], doc: Document):
    in_references = False
    title_done = False

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        i += 1
        if not line:
            continue

        if is_markdown_table_line(line):
            rows: list[list[str]] = []
            while True:
                cells = parse_table_row(line)
                if not is_separator_row(cells):
                    rows.append(cells)
                if i >= len(lines) or not is_markdown_table_line(lines[i]):
                    break
                line = lines[i].strip()
                i += 1
            add_markdown_table(doc, rows)
            continue

        if line.startswith("# "):
            if not title_done:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(line[2:].strip())
                set_run_font(run, 16, True, "000000")
                title_done = True
            continue

        if line == "## 摘要":
            p = doc.add_heading("摘要", level=1)
            continue
        if line == "## 关键词":
            p = doc.add_heading("关键词", level=1)
            continue
        if line == "## 参考文献":
            doc.add_heading("参考文献", level=1)
            in_references = True
            continue

        if line.startswith("## "):
            in_references = False
            doc.add_heading(line[3:].strip(), level=1)
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue

        if in_references and re.match(r"^\[[0-9]+\]", line):
            add_reference_paragraph(doc, line)
        elif line.startswith("工程类学科竞赛；"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(line)
            set_run_font(run, 11)
        elif line.startswith("图1 ") or line.startswith("表"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(line)
            set_run_font(run, 10, False, "555555")
        else:
            add_body_paragraph(doc, line)


def add_metadata(doc: Document):
    props = doc.core_properties
    props.title = "高校工程类学科竞赛赛教融合育人模式构建与实践研究"
    props.subject = "中文省刊投稿初稿"
    props.author = ""
    props.comments = "投稿前请按目标期刊要求补充作者、单位、基金项目等信息。"


def main():
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    parse_markdown(text.splitlines(), doc)
    add_metadata(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
